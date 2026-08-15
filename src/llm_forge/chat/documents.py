"""Document reading with a vision fallback.

Plain text extraction is not enough for a folder of real documents. In a
typical set of 15 PDFs, three were image-only -- a scanned page and two
certificates -- and returned *zero* characters. A reader that only extracts
text reports success and hands back nothing, which is worse than failing:
the model believes it has read the file.

So each page is extracted as text, and any page whose text is too thin to be
real content is rendered to an image and transcribed by a vision-capable
model. Embedded figures can be described the same way.

Requires ``pymupdf`` for PDFs. Vision requires a model that accepts images --
:func:`llm_forge.chat.ollama_provider.supports_vision` checks that.
"""

from __future__ import annotations

import base64
import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

# Below this many characters, a page is almost certainly a scan or a figure
# rather than extractable text. A near-empty page of real text is rare; a
# certificate rendered as a single image is not.
TEXT_THRESHOLD_CHARS = 100

# Rendering above ~150 DPI mostly grows the payload without helping the model.
RENDER_DPI = 150

# Hard ceiling on the rendered long edge. A high-resolution scan at 150 DPI can
# reach 3686x5219, which encodes to a 23 MB PNG -- the API rejects that body
# with a bare 400 ("failed to read request body"), which reads like a model
# problem rather than a payload problem. Downscale before it becomes one.
MAX_IMAGE_EDGE_PX = 2000

# Above this, re-encode as JPEG. Scans compress far better as JPEG than PNG,
# and text stays legible at high quality.
JPEG_THRESHOLD_BYTES = 4 * 1024 * 1024

# Guard against a 200-page PDF becoming 200 vision calls by accident.
DEFAULT_MAX_VISION_PAGES = 10

# Vision transcription is non-deterministic on dense scans; an empty reply is
# often transient rather than a genuinely blank page.
VISION_ATTEMPTS = 3


class DocumentError(RuntimeError):
    """Reading a document failed, with a message safe to show the user."""


@dataclass
class PageResult:
    """One page's extracted content and how it was obtained."""

    number: int
    text: str
    source: str  # "text" | "vision" | "empty"
    note: str = ""

    def as_dict(self) -> dict[str, Any]:
        out: dict[str, Any] = {
            "page": self.number,
            "source": self.source,
            "chars": len(self.text),
            "text": self.text,
        }
        if self.note:
            out["note"] = self.note
        return out


@dataclass
class DocumentResult:
    """A whole document's content, page by page."""

    path: str
    page_count: int
    pages: list[PageResult] = field(default_factory=list)
    images_analyzed: list[dict[str, Any]] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n\n".join(f"--- page {p.number} ---\n{p.text}" for p in self.pages if p.text)

    def as_dict(self) -> dict[str, Any]:
        vision_pages = [p.number for p in self.pages if p.source == "vision"]
        return {
            "status": "ok",
            "path": self.path,
            "page_count": self.page_count,
            "total_chars": len(self.text),
            "pages_needing_vision": vision_pages,
            "pages": [p.as_dict() for p in self.pages],
            "images_analyzed": self.images_analyzed,
            "text": self.text,
        }


def _encode_png(data: bytes) -> str:
    return base64.b64encode(data).decode()


def render_page_image(page: Any, dpi: int = RENDER_DPI) -> tuple[str, str]:
    """Render a PDF page to a base64 image sized to survive the API.

    Returns ``(base64_data, media_type)``. Scales down so the long edge stays
    within :data:`MAX_IMAGE_EDGE_PX`, then re-encodes large results as JPEG.
    """
    pixmap = page.get_pixmap(dpi=dpi)

    longest = max(pixmap.width, pixmap.height)
    if longest > MAX_IMAGE_EDGE_PX:
        # Re-render rather than resample: cheaper and sharper for text.
        scaled_dpi = max(36, int(dpi * MAX_IMAGE_EDGE_PX / longest))
        pixmap = page.get_pixmap(dpi=scaled_dpi)

    data = pixmap.tobytes("png")
    media_type = "image/png"

    if len(data) > JPEG_THRESHOLD_BYTES:
        try:
            data = pixmap.tobytes("jpg", jpg_quality=85)
            media_type = "image/jpeg"
        except Exception:
            # Older PyMuPDF builds without JPEG support: keep the PNG.
            pass

    return _encode_png(data), media_type


def describe_image(
    image_b64: str,
    prompt: str,
    model: str | None = None,
    media_type: str = "image/png",
    client: Any = None,
) -> str:
    """Send one image to a vision-capable model and return its description."""
    from llm_forge.chat.ollama_provider import OllamaError, _client, default_vision_model

    if client is None:
        client = _client()
    if model is None:
        model = default_vision_model(client=client)
    if not model:
        raise DocumentError(
            "No vision-capable model is available. Run /model to see what your "
            "key can reach; kimi-k2.6 and gemma4:31b both accept images."
        )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{media_type};base64,{image_b64}"},
                        },
                    ],
                }
            ],
            max_tokens=4096,
        )
    except Exception as exc:
        from llm_forge.chat.ollama_provider import _explain

        raise DocumentError(_explain(exc)) from exc

    return (response.choices[0].message.content or "").strip()


TRANSCRIBE_PROMPT = (
    "Transcribe all text visible in this page image, preserving the reading "
    "order and any table structure. If it is a certificate, form, or ID, "
    "include every field and its value. Reply with the transcription only -- "
    "no preamble, no commentary. If the page has no legible text, reply "
    "exactly: [no legible text]"
)

DESCRIBE_FIGURE_PROMPT = (
    "Describe this figure from a document: what it shows, its type (chart, "
    "diagram, photo, logo, signature), and any text or numbers it contains. "
    "Be specific and concise."
)


def read_pdf(
    path: Path,
    use_vision: bool = True,
    max_vision_pages: int = DEFAULT_MAX_VISION_PAGES,
    analyze_figures: bool = False,
    model: str | None = None,
    client: Any = None,
) -> DocumentResult:
    """Read a PDF page by page, using vision where text extraction comes up empty."""
    try:
        import pymupdf
    except ImportError:
        try:
            import fitz as pymupdf  # older releases
        except ImportError as exc:
            raise DocumentError(
                "Reading PDFs needs pymupdf. Install it with: pip install pymupdf"
            ) from exc

    try:
        doc = pymupdf.open(str(path))
    except Exception as exc:
        raise DocumentError(f"Could not open {path.name}: {exc}") from exc

    result = DocumentResult(path=str(path), page_count=len(doc))
    vision_used = 0

    try:
        for number, page in enumerate(doc, start=1):
            text = (page.get_text() or "").strip()

            if len(text) >= TEXT_THRESHOLD_CHARS:
                result.pages.append(PageResult(number, text, "text"))
                continue

            if not use_vision:
                result.pages.append(
                    PageResult(
                        number,
                        text,
                        "empty" if not text else "text",
                        note="Little or no extractable text; vision is disabled.",
                    )
                )
                continue

            if vision_used >= max_vision_pages:
                result.pages.append(
                    PageResult(
                        number,
                        text,
                        "empty",
                        note=(
                            f"Skipped vision: max_vision_pages={max_vision_pages} "
                            "reached. Raise it to transcribe this page."
                        ),
                    )
                )
                continue

            # Scanned or image-only page -- render it and let a vision model read it.
            try:
                image_b64, media_type = render_page_image(page)
                # Vision calls on dense scans intermittently come back empty --
                # the same page transcribed 914 characters on a retry after
                # returning nothing. Retry before concluding the page is blank.
                transcript = ""
                for attempt in range(VISION_ATTEMPTS):
                    transcript = describe_image(
                        image_b64,
                        TRANSCRIBE_PROMPT,
                        model=model,
                        media_type=media_type,
                        client=client,
                    )
                    if transcript.strip():
                        break
                vision_used += 1
                # An empty reply is a failed read, not a successful one. Marking
                # it "vision" made the page look transcribed while carrying no
                # text -- a silent zero that is worse than a reported failure.
                if not transcript.strip() or transcript.strip() == "[no legible text]":
                    result.pages.append(
                        PageResult(
                            number,
                            "",
                            "empty",
                            note=(
                                f"Vision returned no text after {VISION_ATTEMPTS} "
                                "attempts -- the page may be blank or too "
                                "low-contrast to read."
                            ),
                        )
                    )
                else:
                    result.pages.append(
                        PageResult(
                            number,
                            transcript,
                            "vision",
                            note="Page had no extractable text; transcribed from an image.",
                        )
                    )
            except DocumentError as exc:
                result.pages.append(
                    PageResult(number, text, "empty", note=f"Vision failed: {exc}")
                )

        if analyze_figures:
            result.images_analyzed = _analyze_figures(
                doc, pymupdf, model=model, client=client, limit=max_vision_pages
            )
    finally:
        doc.close()

    return result


def _analyze_figures(
    doc: Any, pymupdf: Any, model: str | None, client: Any, limit: int
) -> list[dict[str, Any]]:
    """Describe images embedded in the document."""
    described: list[dict[str, Any]] = []

    for page_number, page in enumerate(doc, start=1):
        for image in page.get_images(full=True):
            if len(described) >= limit:
                return described
            xref = image[0]
            try:
                pixmap = pymupdf.Pixmap(doc, xref)
                # CMYK and alpha variants need converting before PNG encoding.
                if pixmap.n - pixmap.alpha >= 4:
                    pixmap = pymupdf.Pixmap(pymupdf.csRGB, pixmap)
                # Icons and rules are not worth a vision call.
                if pixmap.width < 100 or pixmap.height < 100:
                    continue
                description = describe_image(
                    _encode_png(pixmap.tobytes("png")),
                    DESCRIBE_FIGURE_PROMPT,
                    model=model,
                    client=client,
                )
                described.append(
                    {"page": page_number, "xref": xref, "description": description}
                )
            except Exception as exc:
                described.append({"page": page_number, "xref": xref, "error": str(exc)})

    return described


def read_docx(path: Path) -> DocumentResult:
    """Read a .docx file's paragraphs and tables."""
    try:
        import docx
    except ImportError as exc:
        raise DocumentError(
            "Reading .docx needs python-docx. Install it with: pip install python-docx"
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocumentError(f"Could not open {path.name}: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return DocumentResult(
        path=str(path),
        page_count=1,
        pages=[PageResult(1, text, "text" if text else "empty")],
    )


def read_document(
    path: str,
    use_vision: bool = True,
    analyze_figures: bool = False,
    max_vision_pages: int = DEFAULT_MAX_VISION_PAGES,
    model: str | None = None,
    client: Any = None,
) -> DocumentResult:
    """Read any supported document, falling back to vision for image-only pages."""
    file_path = Path(path).expanduser().resolve()

    if not file_path.exists():
        raise DocumentError(f"File not found: {path}")
    if not file_path.is_file():
        raise DocumentError(f"Not a file: {path}")

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return read_pdf(
            file_path,
            use_vision=use_vision,
            max_vision_pages=max_vision_pages,
            analyze_figures=analyze_figures,
            model=model,
            client=client,
        )

    if suffix == ".docx":
        return read_docx(file_path)

    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}:
        media = "image/jpeg" if suffix in {".jpg", ".jpeg"} else f"image/{suffix.lstrip('.')}"
        transcript = describe_image(
            _encode_png(file_path.read_bytes()),
            TRANSCRIBE_PROMPT,
            model=model,
            media_type=media,
            client=client,
        )
        return DocumentResult(
            path=str(file_path),
            page_count=1,
            pages=[PageResult(1, transcript, "vision")],
        )

    if suffix in {".txt", ".md", ".rst", ".csv", ".json", ".jsonl", ".yaml", ".yml"}:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return DocumentResult(
            path=str(file_path), page_count=1, pages=[PageResult(1, text, "text")]
        )

    raise DocumentError(
        f"Unsupported file type '{suffix}'. Supported: .pdf, .docx, images, "
        "and plain-text formats."
    )


def read_folder(
    folder: str,
    use_vision: bool = True,
    max_files: int = 25,
    max_vision_pages: int = DEFAULT_MAX_VISION_PAGES,
    model: str | None = None,
    client: Any = None,
) -> dict[str, Any]:
    """Read every supported document in a folder.

    This is the operation the assistant actually needs when a user says "read
    my folder" -- doing it file by file costs a round trip each and the model
    has to know the filenames first.
    """
    directory = Path(folder).expanduser().resolve()
    if not directory.is_dir():
        raise DocumentError(f"Not a directory: {folder}")

    supported = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".jsonl"}
    candidates = sorted(
        p for p in directory.iterdir() if p.is_file() and p.suffix.lower() in supported
    )

    documents: list[dict[str, Any]] = []
    failures: list[dict[str, str]] = []

    for file_path in candidates[:max_files]:
        try:
            result = read_document(
                str(file_path),
                use_vision=use_vision,
                max_vision_pages=max_vision_pages,
                model=model,
                client=client,
            )
            documents.append(
                {
                    "file": file_path.name,
                    "page_count": result.page_count,
                    "chars": len(result.text),
                    "pages_needing_vision": [
                        p.number for p in result.pages if p.source == "vision"
                    ],
                    # Without this a zero-character file gives no reason why,
                    # and the model cannot tell "blank scan" from "read failed".
                    "warnings": [
                        f"page {p.number}: {p.note}" for p in result.pages if p.note
                    ],
                    "text": result.text,
                }
            )
        except DocumentError as exc:
            failures.append({"file": file_path.name, "error": str(exc)})

    return {
        "status": "ok",
        "folder": str(directory),
        "files_found": len(candidates),
        "files_read": len(documents),
        "truncated": len(candidates) > max_files,
        "documents": documents,
        "failures": failures,
        "total_chars": sum(d["chars"] for d in documents),
    }


def read_document_json(**kwargs: Any) -> str:
    """JSON-returning wrapper for the tool dispatcher."""
    try:
        return json.dumps(read_document(**kwargs).as_dict(), indent=2, default=str)
    except DocumentError as exc:
        return json.dumps({"status": "error", "error": str(exc)})


def read_folder_json(**kwargs: Any) -> str:
    """JSON-returning wrapper for the tool dispatcher."""
    try:
        return json.dumps(read_folder(**kwargs), indent=2, default=str)
    except DocumentError as exc:
        return json.dumps({"status": "error", "error": str(exc)})


__all__ = [
    "DEFAULT_MAX_VISION_PAGES",
    "DocumentError",
    "DocumentResult",
    "PageResult",
    "TEXT_THRESHOLD_CHARS",
    "describe_image",
    "read_document",
    "read_document_json",
    "read_folder",
    "read_folder_json",
    "read_pdf",
]
